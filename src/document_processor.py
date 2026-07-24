import docx
from docx.text.paragraph import Paragraph
from docx.oxml.ns import qn

from typing import List, Any, Iterator
from .analyzer import RedactionAnalyzer
from .anonymizer import FakeDataAnonymizer

class DocxProcessor:
    """Processes docx files to redact PII while preserving formatting."""
    
    def __init__(self, analyzer: RedactionAnalyzer, anonymizer: FakeDataAnonymizer):
        self.analyzer = analyzer
        self.anonymizer = anonymizer

    def _iter_all_paragraphs(self, doc: docx.Document) -> Iterator[Paragraph]:
        """
        Yield every paragraph in document order, including those inside tables
        and nested tables. python-docx's doc.paragraphs misses table-cell paragraphs;
        doc.tables misses nested tables. This walks the raw XML body to get everything.
        """
        body = doc.element.body
        for child in body.iter():
            if child.tag == qn('w:p'):
                yield Paragraph(child, doc)

    def process_document(self, input_path: str, output_path: str):
        """Main method to process the entire document."""
        doc = docx.Document(input_path)

        # Walk every paragraph in document order (body + all table cells, including nested)
        for para in self._iter_all_paragraphs(doc):
            self._process_paragraph(para)

        doc.save(output_path)

    def _process_paragraph(self, paragraph: Paragraph):
        """Processes a single paragraph, replacing PII at the run level."""
        text = paragraph.text
        if not text.strip():
            return
            
        # Analyze the full paragraph text
        results = self.analyzer.analyze_text(text)
        if not results:
            return
            
        # Sort results by score desc, then span length desc so the best match wins
        # when spans overlap.
        results = sorted(results, key=lambda x: (x.score, x.end - x.start), reverse=True)

        # Deduplicate: remove any result whose span overlaps with a higher-priority
        # result already kept. This prevents double-replacement corruption when
        # Presidio returns both PERSON "ksh" and EMAIL "ksh@domain.com" at the same offset.
        deduped = []
        for res in results:
            if not any(
                res.start < kept.end and res.end > kept.start
                for kept in deduped
            ):
                deduped.append(res)

        # Now sort in reverse start order for safe in-place replacement
        results = sorted(deduped, key=lambda x: x.start, reverse=True)
        
        # We need a robust way to replace text that might span multiple runs.
        # But for simplicity and safety, let's build a map of character indices to runs.
        # Actually, since we are doing replacements, it's easier to find the exact string
        # in the runs, but what if it spans runs?
        
        # Let's map each character index in the paragraph to a (run_index, char_index_in_run)
        run_map = []
        for r_idx, run in enumerate(paragraph.runs):
            for c_idx, char in enumerate(run.text):
                run_map.append((r_idx, c_idx))
                
        # If the text length doesn't match the run_map length, something is weird (e.g. fields)
        if len(text) != len(run_map):
            # Fallback: simple string replace in each run
            self._fallback_process(paragraph, results, text)
            return

        # Perform replacement
        # Since we iterate in reverse, replacing characters won't affect the indices of earlier results
        for res in results:
            original_text = text[res.start:res.end]
            fake_text = self.anonymizer.get_fake_value(res.entity_type, original_text)
            
            # The runs involved in this span
            span_runs = run_map[res.start:res.end]
            if not span_runs:
                continue
                
            # Group by run_index
            runs_involved = {}
            for r_idx, c_idx in span_runs:
                if r_idx not in runs_involved:
                    runs_involved[r_idx] = []
                runs_involved[r_idx].append(c_idx)
                
            # Replace: put the fake text in the first run involved, and delete characters from all involved runs
            first_run_idx = list(runs_involved.keys())[0]
            
            for r_idx, char_indices in runs_involved.items():
                run = paragraph.runs[r_idx]
                run_text_list = list(run.text)
                
                # Delete the characters that are part of the PII
                # We sort reverse so deleting doesn't shift indices of the same run
                for c_idx in sorted(char_indices, reverse=True):
                    if c_idx < len(run_text_list):
                        run_text_list.pop(c_idx)
                
                # Insert the fake text at the position of the first deleted character in the first run
                if r_idx == first_run_idx:
                    min_idx = min(char_indices)
                    run_text_list.insert(min_idx, fake_text)
                    
                run.text = "".join(run_text_list)
                
            # Rebuild the run_map and text because we modified the runs
            # Actually, since we process in reverse, we don't strictly need to rebuild IF we only
            # rely on indices before `res.start`.
            # But let's rebuild just in case to be safe, though it's computationally heavier.
            # Wait, if we rebuild, the indices for the remaining results (which are earlier in the text)
            # would STILL be valid because they are before the replaced text!
            # So we DON'T rebuild run_map here.

    def _fallback_process(self, paragraph: Paragraph, results: List[Any], text: str):
        """Simple fallback that replaces text in runs if exact match is found."""
        for res in results:
            original_text = text[res.start:res.end]
            fake_text = self.anonymizer.get_fake_value(res.entity_type, original_text)
            for run in paragraph.runs:
                if original_text in run.text:
                    run.text = run.text.replace(original_text, fake_text)
